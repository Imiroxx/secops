"""
Скрипт форматирования научной статьи по требованиям
Требования из файла: Требования к работам_4.pdf
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_format(paragraph, font_name='Times New Roman', font_size=14, 
                         bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(1.25), line_spacing=1.5,
                         space_before=0, space_after=0):
    """Устанавливает форматирование абзаца"""
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = first_line_indent
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_custom(doc, text, level=1):
    """Добавляет заголовок с заданным форматированием"""
    if level == 0:  # Заголовок работы
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)
    elif level == 1:  # Раздел (ГЛАВА 1)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)
    elif level == 2:  # Подраздел (1.1)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
    else:  # Подподраздел (1.1.1)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
    return p


def add_text_paragraph(doc, text, first_indent=True):
    """Добавляет текстовый абзац с отступом"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def create_title_page(doc):
    """Создание титульной страницы"""
    # Шапка
    p = doc.add_paragraph()
    run = p.add_run('МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И МАССОВЫХ КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('«СИБИРСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ТЕЛЕКОММУНИКАЦИЙ И ИНФОРМАТИКИ»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('Факультет информационных систем и технологий')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('Кафедра защиты информации')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Отступы
    for _ in range(3):
        doc.add_paragraph()
    
    # Название работы
    p = doc.add_paragraph()
    run = p.add_run('СИММЕТРИЧНЫЕ И АСИММЕТРИЧНЫЕ АЛГОРИТМЫ ШИФРОВАНИЯ В СОВРЕМЕННЫХ ИНФОРМАЦИОННЫХ СИСТЕМАХ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Пояснительная записка к научной работе')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('по дисциплине: «Информационная безопасность»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Отступы
    for _ in range(4):
        doc.add_paragraph()
    
    # Автор и проверяющий
    p = doc.add_paragraph()
    run = p.add_run('Выполнил: студент группы ИС-123')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph()
    run = p.add_run('Ямалетдинов К.Р.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Проверил: доцент кафедры ЗИ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph()
    run = p.add_run('Иванов И.И.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Отступ
    for _ in range(6):
        doc.add_paragraph()
    
    # Место и год
    p = doc.add_paragraph()
    run = p.add_run('Новосибирск 2026')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_table_of_contents(doc):
    """Создание содержания"""
    p = doc.add_paragraph()
    run = p.add_run('СОДЕРЖАНИЕ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    toc_items = [
        ('Введение', 3),
        ('1 Теоретические основы криптографической защиты информации', 5),
        ('1.1 Основные понятия и определения криптографии', 5),
        ('1.2 Классификация методов шифрования', 7),
        ('1.3 Требования к современным криптографическим системам', 9),
        ('2 Симметричные алгоритмы шифрования', 11),
        ('2.1 Алгоритм AES (Advanced Encryption Standard)', 11),
        ('2.1.1 Структура алгоритма AES', 12),
        ('2.1.2 Операции преобразования данных в AES', 14),
        ('2.2 Алгоритм Triple DES', 16),
        ('2.2.1 Принцип работы Triple DES', 16),
        ('2.2.2 Сравнительный анализ AES и Triple DES', 18),
        ('3 Асимметричные алгоритмы шифрования', 20),
        ('3.1 Алгоритм RSA', 20),
        ('3.1.1 Математические основы RSA', 21),
        ('3.1.2 Генерация ключей и процесс шифрования', 23),
        ('3.2 Алгоритм ECC (Elliptic Curve Cryptography)', 25),
        ('3.2.1 Особенности эллиптических кривых', 26),
        ('3.2.2 Преимущества ECC перед RSA', 28),
        ('4 Программная реализация криптографических алгоритмов', 30),
        ('4.1 Реализация симметричного шифрования на Python', 30),
        ('4.2 Реализация асимметричного шифрования на Python', 33),
        ('4.3 Сравнительный анализ производительности', 35),
        ('Заключение', 37),
        ('Список использованной литературы', 39),
        ('Приложение А Листинг программного кода', 41),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Добавляем точки и номер страницы
        tab_run = p.add_run('\t' + str(page))
        tab_run.font.name = 'Times New Roman'
        tab_run.font.size = Pt(14)
        tab_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)


def create_introduction(doc):
    """Создание введения"""
    add_heading_custom(doc, 'ВВЕДЕНИЕ', level=1)
    
    intro_texts = [
        'В современном информационном обществе проблема защиты конфиденциальной информации приобретает особую актуальность. Ежедневно в глобальной сети Интернет передаются огромные объемы данных, содержащих персональную, коммерческую и государственную тайну. Развитие технологий квантовых вычислений, рост вычислительной мощности компьютеров и совершенствование методов криптоанализа диктуют необходимость постоянного совершенствования криптографических систем защиты информации.',
        'Актуальность выбранной темы обусловлена несколькими факторами. Во-первых, увеличение числа кибератак на критическую информационную инфраструктуру требует применения надежных методов шифрования. Во-вторых, появление квантовых компьютеров ставит под угрозу классические асимметричные алгоритмы, что стимулирует разработку постквантовых криптографических систем. В-третьих, рост объемов передаваемых данных в условиях ограниченных вычислительных ресурсов мобильных устройств требует оптимизации криптографических алгоритмов.',
        'Целью данной работы является комплексный анализ симметричных и асимметричных алгоритмов шифрования, применяемых в современных информационных системах, и разработка программной реализации наиболее распространенных криптографических методов.',
        'Для достижения поставленной цели необходимо решить следующие задачи:',
        'изучить теоретические основы криптографической защиты информации;',
        'провести анализ симметричных алгоритмов шифрования на примере AES и Triple DES;',
        'исследовать асимметричные алгоритмы шифрования, включая RSA и ECC;',
        'разработать программную реализацию криптографических алгоритмов на языке Python;',
        'провести сравнительный анализ производительности реализованных алгоритмов.',
        'Объектом исследования являются процессы шифрования и дешифрования данных в информационных системах.',
        'Предметом исследования выступают симметричные и асимметричные алгоритмы криптографической защиты информации.',
        'Методологическую основу исследования составляют системный анализ, сравнительный анализ, теоретическое моделирование и эмпирические методы исследования.',
        'Теоретическая значимость работы заключается в обобщении и систематизации знаний о современных криптографических алгоритмах, а также в разработке методических рекомендаций по выбору оптимальных алгоритмов шифрования для различных сценариев применения.',
        'Практическая ценность полученных результатов состоит в создании программного комплекса, реализующего основные криптографические алгоритмы, который может быть использован в учебном процессе и при разработке систем защиты информации.',
        'В работе использованы научные труды отечественных и зарубежных авторов в области криптографии и информационной безопасности, в том числе работы А.А. Альбертина, Б.А. Молдовян, Д. Стинсона, Б. Шнайера, а также документация стандартов FIPS и RFC.',
    ]
    
    for i, text in enumerate(intro_texts):
        if i == 3:  # Перед списком задач
            add_text_paragraph(doc, text)
        elif 4 <= i <= 8:  # Список задач
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(1.25)
        else:
            add_text_paragraph(doc, text)


def format_article():
    """Основная функция форматирования статьи"""
    doc = Document()
    
    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # Титульная страница
    create_title_page(doc)
    
    # Разрыв страницы
    doc.add_page_break()
    
    # Содержание
    create_table_of_contents(doc)
    
    # Разрыв страницы
    doc.add_page_break()
    
    # Введение
    create_introduction(doc)
    
    # Сохранение документа
    doc.save('article_formatted.docx')
    print("Форматированная статья сохранена в файл: article_formatted.docx")


if __name__ == '__main__':
    format_article()
