#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор обновлённых тезисов Ямалетдинова Карима по новому положению
Требования: 1 страница, Times New Roman 12пт, интервал 1, поля 2 см, отступ 1.25 см
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis():
    doc = Document()
    
    # Настройка полей: 2 см со всех сторон
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Заголовок (по центру, полужирный, 12пт)
    title = doc.add_paragraph()
    run = title.add_run("TrustProject – веб-приложение для комплексного анализа безопасности URL-адресов и файлов")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    
    # Автор (по центру, обычный, 12пт)
    author = doc.add_paragraph()
    run = author.add_run("Ямалетдинов Карим, 10 класс")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(0)
    
    # Организация (по центру, курсив, 12пт)
    org = doc.add_paragraph()
    run = org.add_run("МБОУ СОШ № ___, г. Казань")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(0)
    
    # Научный руководитель (по центру, курсив, 12пт)
    supervisor = doc.add_paragraph()
    run = supervisor.add_run("Научный руководитель: Петрова О.В.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor.paragraph_format.space_after = Pt(12)
    
    # Функция для добавления абзаца с форматированием
    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        return p
    
    # Введение / актуальность
    add_para("Современное интернет-пространство динамично развивается: более 2 миллиардов веб-сайтов и 500 тысяч новых файлов ежедневно. По данным Positive Technologies, количество кибератак увеличилось на 32% в 2024 году, 68% атак направлены на рядовых пользователей. В образовательной среде ситуация особенно критична: 74% школьников сталкивались с мошенническими сайтами, 56% загружали вредоносные файлы.")
    
    add_para("Цель работы: разработка веб-приложения TrustProject – комплексной системы для анализа безопасности URL-адресов и файлов, интегрирующей современные технологии детектирования угроз.", bold=True)
    
    add_para("Задачи: провести аналитический обзор существующих решений для проверки безопасности; изучить теоретические основы детектирования вредоносного контента; разработать архитектуру и реализовать систему TrustProject; провести тестирование системы.", bold=True)
    
    add_para("Проведён аналитический обзор предметной области. Профессиональные платформы (VirusTotal, Hybrid Analysis) обладают высокой точностью, но сложны для неподготовленных пользователей и имеют ограничения бесплатного доступа. Браузерные расширения (Norton Safe Web) удобны, но ограничены функциональностью. Онлайн-сервисы (Sucuri) просты, но имеют недостаточную глубину анализа.")
    
    add_para("Изучены теоретические основы анализа безопасности: статический анализ (проверка без выполнения), динамический анализ (в песочнице), гибридный анализ (комбинация методов). Применяются энтропийный анализ (вычисление энтропии Шеннона), сигнатурный анализ, проверка SSL/TLS сертификатов.")
    
    add_para("Разработана архитектура системы с трёхзвенной структурой: уровень представления (HTML5/CSS3/JavaScript), уровень бизнес-логики (Python, Flask), уровень данных (SQLite). Реализованы функциональные модули: анализ URL (интеграция с VirusTotal API, Google Safe Browsing API, WHOIS, AbuseIPDB), анализ файлов (SHA-256, статический и поведенческий анализ, Hybrid Analysis API), социальный модуль (комментарии, голосования), визуализация (Chart.js, vis-network).")
    
    add_para("Система оценки доверия рассчитывает общую оценку безопасности от 0 до 100 баллов по формуле: Total_Score = (AV_Score × 0.4) + (Behavior_Score × 0.25) + (Reputation_Score × 0.15) + (Technical_Score × 0.1) + (Social_Score × 0.1). Цветовая кодировка: красный (0-39, опасный), оранжевый (40-69, подозрительный), жёлтый (70-89, умеренно безопасный), зелёный (90-100, безопасный).")
    
    add_para("Технологический стек: Python 3.12+, Flask 2.3.3, SQLAlchemy (backend); HTML5, CSS3, Tailwind CSS, Chart.js (frontend); SQLite (база данных); интеграции с VirusTotal API, Google Safe Browsing API, Hybrid Analysis API, AbuseIPDB API.")
    
    add_para("Преимущества разработанного приложения: комплексность – единый интерфейс для URL и файлов; доступность – бесплатное использование без ограничений; понятность – оценка 0-100 с цветовой индикацией; образовательный компонент – система комментариев и голосований; русскоязычный интерфейс.")
    
    add_para("Практическая значимость: инструмент может использоваться образовательными учреждениями для обучения цифровой грамотности, индивидуальными пользователями для проверки подозрительных ссылок и файлов, системными администраторами для мониторинга безопасности.")
    
    doc.save('Ямалетдинов_Карим_тезисы_новое_положение.docx')
    print('Тезисы сохранены: Ямалетдинов_Карим_тезисы_новое_положение.docx')

if __name__ == '__main__':
    create_thesis()
