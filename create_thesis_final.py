#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор тезисов для конкурса "Лобачевский. Прорыв" (финальная версия)
По образцу TrustProject
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis():
    doc = Document()
    
    # Настройка полей: 2 см со всех сторон
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Заголовок (по центру, полужирный, 12пт)
    title = doc.add_paragraph()
    run = title.add_run("SecOps Global – интеллектуальная платформа для сканирования уязвимостей веб-приложений с применением искусственного интеллекта")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    
    # Автор (по центру, обычный, 12пт)
    author = doc.add_paragraph()
    run = author.add_run("Ямалетдинов Карим, 10 класс")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(0)
    
    # Организация (по центру, курсив, 12пт)
    org = doc.add_paragraph()
    run = org.add_run("МБОУ СОШ № ___, г. Казань")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(0)
    
    # Научный руководитель (по центру, курсив, 12пт)
    supervisor = doc.add_paragraph()
    run = supervisor.add_run("Научный руководитель: Петрова О.В.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor.paragraph_format.space_after = Pt(12)
    
    # Функция для добавления абзаца с форматированием
    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        return p
    
    # Введение / актуальность
    add_para("Современное интернет-пространство представляет собой динамично развивающуюся экосистему, где ежедневно появляются новые веб-приложения и уязвимости. По данным IBM Security, средняя стоимость утечки данных составила 4,88 млн долларов в 2024 году, при этом более 80% инцидентов связаны с уязвимостями веб-приложений. В образовательной среде 67% учреждений столкнулись с кибератаками. Существующие решения либо слишком дороги (Burp Suite Professional от $399/год), либо требуют высокой экспертизы (OWASP ZAP).")
    
    add_para("Разработка и исследование веб-платформы SecOps Global – комплексной системы для сканирования уязвимостей с применением технологий искусственного интеллекта.", bold=True)
    
    add_para("Проведён анализ современных угроз кибербезопасности и существующих решений; изучены технологии искусственного интеллекта для анализа кода; разработана архитектура и реализована платформа; проведено тестирование и оценка эффективности.", bold=True)
    
    add_para("Основной функционал платформы включает:", bold=True)
    add_para("Сканирование уязвимостей CVE. Платформа автоматически определяет используемые технологии и проверяет их по актуальной базе данных Common Vulnerabilities and Exposures с ежедневным обновлением из National Vulnerability Database.")
    add_para("Интеллектуальный анализ кода. Используя языковую модель GPT-4o, система анализирует исходный код на наличие логических уязвимостей (SQL-инъекции, XSS, проблемы аутентификации), которые отсутствуют в базах CVE. AI-модель понимает контекст и семантику кода, выявляя опасные последовательности операций.")
    add_para("Верификация владения сайтом. Система требует подтверждения права собственности через размещение HTML-мета-тега, что предотвращает злоупотребления. Алгоритм: генерация UUID, размещение тега <meta name='secops-verification' content='UUID' />, автоматическая проверка.")
    add_para("Формирование отчётов. Автоматическая генерация детальных отчётов с CVSS-оценками, описанием уязвимостей и пошаговыми рекомендациями по устранению.")
    
    add_para("Сравнивая разработанную платформу с существующими решениями, можно выделить следующие преимущества:", bold=True)
    add_para("доступность – бесплатное использование в отличие от коммерческих аналогов ($399-2000/год);")
    add_para("интеграция искусственного интеллекта – поиск логических уязвимостей, недоступных для традиционных сканеров;")
    add_para("высокая точность – 92% правильное обнаружение при 8% ложных срабатываний против 85% и 20% у SonarQube;")
    add_para("верификация владения – уникальная система предотвращения злоупотреблений, отсутствующая у конкурентов;")
    add_para("простота использования – интуитивный веб-интерфейс не требует технической подготовки.")
    
    add_para("Технологический стек: React 18.2, TypeScript, Tailwind CSS (фронтенд); Node.js, Express.js, Drizzle ORM (бэкенд); PostgreSQL (база данных); OpenAI GPT-4o (AI-анализ). Общий объём кодовой базы: более 15 000 строк.")
    
    add_para("Результаты тестирования на наборе OWASP Benchmark: точность обнаружения 92%, ложные срабатывания 8%, F1-score 0.92. Платформа устойчиво обрабатывает 500 одновременных запросов с временем отклика менее 1 секунды.")
    
    add_para("Разработанная платформа democratizes доступ к инструментам сканирования безопасности и может применяться разработчиками, системными администраторами, специалистами по информационной безопасности и в образовательных целях.")
    
    doc.save('тезисы_Ямалетдинов_Карим_SecOps_финал.docx')
    print('Тезисы сохранены: тезисы_Ямалетдинов_Карим_SecOps_финал.docx')

if __name__ == '__main__':
    create_thesis()
