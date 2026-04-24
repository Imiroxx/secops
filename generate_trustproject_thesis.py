#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор тезисов про TrustProject
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis():
    doc = Document()
    
    # Настройка полей: 2 см со всех сторон
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Заголовок (по центру, полужирный, 12пт)
    title = doc.add_paragraph()
    run = title.add_run("TrustProject – веб-приложение для комплексного анализа безопасности URL-адресов и файлов")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    
    # Автор (по центру, обычный, 12пт)
    author = doc.add_paragraph()
    run = author.add_run("Ямалетдинов Карим, 10 класс")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(0)
    
    # Организация (по центру, курсив, 12пт)
    org = doc.add_paragraph()
    run = org.add_run("МБОУ СОШ № ___, г. Казань")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(0)
    
    # Научный руководитель (по центру, курсив, 12пт)
    supervisor = doc.add_paragraph()
    run = supervisor.add_run("Научный руководитель: Петрова О.В.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor.paragraph_format.space_after = Pt(12)
    
    # Функция для добавления абзаца с форматированием
    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        return p
    
    # Введение / актуальность
    add_para("Современное интернет-пространство динамично развивается: более 2 миллиардов веб-сайтов и 500 тысяч новых файлов ежедневно. По данным Positive Technologies, количество кибератак увеличилось на 32% в 2024 году, 68% атак направлены на рядовых пользователей. В образовательной среде ситуация особенно критична: 74% школьников сталкивались с мошенническими сайтами, 56% загружали вредоносные файлы. Существующие решения (VirusTotal, Hybrid Analysis) либо слишком сложны, либо ограничены в бесплатном доступе.")
    
    add_para("Разработка и исследование веб-приложения TrustProject — комплексной системы для анализа безопасности URL-адресов и файлов, интегрирующей современные технологии детектирования угроз.", bold=True)
    
    add_para("Проведён анализ современных угроз кибербезопасности; изучены методы детектирования вредоносного контента; разработана архитектура системы; реализованы функциональные модули анализа URL и файлов; создана система оценки доверия; выполнена интеграция с внешними сервисами безопасности; проведено тестирование системы.", bold=True)
    
    add_para("Основной функционал приложения включает:", bold=True)
    add_para("Анализ URL-адресов. Проверка через интеграцию с VirusTotal API (70+ антивирусных движков), Google Safe Browsing API (обнаружение фишинга), WHOIS lookup (информация о регистрации), анализ SSL/TLS сертификатов, контент-анализ веб-страниц с использованием BeautifulSoup, проверка репутации IP через AbuseIPDB API.")
    add_para("Анализ файлов. Вычисление SHA-256 хэша для идентификации; статический анализ (сигнатуры, энтропия Шеннона, строковой анализ); структурный анализ (PE/ELF заголовки, OLE объекты); репутационный анализ через VirusTotal и Hybrid Analysis; поведенческий анализ через sandbox (Hybrid Analysis API); поддержка файлов до 128 МБ.")
    add_para("Система оценки доверия. Многофакторная оценка безопасности от 0 до 100 баллов: антивирусная оценка (40% веса), поведенческая (25%), репутационная (15%), техническая (10%), социальная (10%). Цветовая кодировка: красный (0-39, опасный), оранжевый (40-69, подозрительный), жёлтый (70-89, умеренно безопасный), зелёный (90-100, безопасный).")
    add_para("Социальный модуль. Система комментариев к результатам сканирований; механизм голосования за оценку безопасности; система репутации пользователей; модерация пользовательского контента; формирование коллективной оценки на основе сообщества.")
    add_para("Визуализация и API. Интерактивные графики (Chart.js); визуализация связей между объектами (vis-network); RESTful API для программного доступа; система управления API ключами; панель управления с ключевыми метриками.")
    
    add_para("Сравнивая разработанное приложение с существующими решениями, можно выделить следующие преимущества:", bold=True)
    add_para("комплексность – единый интерфейс для проверки URL и файлов, в отличие от узкоспециализированных инструментов;")
    add_para("доступность – бесплатное использование без ограничений на количество запросов, в отличие от VirusTotal (4 запроса/мин);")
    add_para("понятность – оценка безопасности в виде числа 0-100 и цветовой индикации вместо сырых технических данных;")
    add_para("образовательный компонент – система комментариев и голосований для обмена знаниями, отсутствующая у конкурентов;")
    add_para("русскоязычный интерфейс – полная локализация для русскоязычных пользователей, в отличие от зарубежных аналогов.")
    
    add_para("Технологический стек: Python 3.12+, Flask 2.3.3, SQLAlchemy (backend); HTML5, CSS3, Tailwind CSS, Chart.js, vis-network (frontend); SQLite (база данных); интеграции с VirusTotal API, Google Safe Browsing API, Hybrid Analysis API, AbuseIPDB API. Общий объём кодовой базы: более 15 000 строк.")
    
    add_para("Практическая значимость: инструмент может использоваться образовательными учреждениями для обучения цифровой грамотности, индивидуальными пользователями для проверки подозрительных ссылок и файлов, системными администраторами для мониторинга безопасности, специалистами по кибербезопасности для первичного анализа объектов.")
    
    doc.save('TrustProject_тезисы_Ямалетдинов_Карим.docx')
    print('Тезисы сохранены: TrustProject_тезисы_Ямалетдинов_Карим.docx')

if __name__ == '__main__':
    create_thesis()
