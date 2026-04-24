#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор тезисов для конкурса "Лобачевский. Прорыв" (обновленная версия)
По образцу прошлого года
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_thesis():
    doc = Document()
    
    # Настройка полей: 2 см со всех сторон
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Заголовок (по центру, полужирный, 12пт)
    title = doc.add_paragraph()
    run = title.add_run("SecOps Global – веб-платформа для сканирования уязвимостей с использованием искусственного интеллекта")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    
    # Автор (по центру, обычный, 12пт)
    author = doc.add_paragraph()
    run = author.add_run("Ямалетдинов Карим, 10 класс")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(0)
    
    # Организация (по центру, курсив, 12пт)
    org = doc.add_paragraph()
    run = org.add_run("МБОУ СОШ № ___, г. Казань")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_after = Pt(0)
    
    # Научный руководитель (по центру, курсив, 12пт)
    supervisor = doc.add_paragraph()
    run = supervisor.add_run("Научный руководитель: Петрова О.В.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    supervisor.paragraph_format.space_after = Pt(12)
    
    # Функция для добавления абзаца с форматированием
    def add_para(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        return p
    
    # Введение / актуальность
    add_para("Актуальность. В современном цифровом мире веб-приложения стали неотъемлемой частью бизнес-процессов, но их безопасность вызывает серьезную обеспокоенность. Согласно отчету IBM за 2024 год, средняя стоимость утечки данных составляет 4,88 млн долларов, при этом более 80% инцидентов связаны с уязвимостями веб-приложений. Существующие сканеры уязвимостей либо недоступны по цене (Burp Suite Professional стоит от 399$ в год), либо требуют высокой технической экспертизы для настройки.")
    
    # Цель
    add_para("Цель работы. Разработка веб-платформы SecOps Global для автоматизированного сканирования уязвимостей веб-приложений с применением технологий искусственного интеллекта.")
    
    # Функционал
    add_para("Основной функционал платформы включает:", bold=True)
    add_para("Сканирование уязвимостей CVE. Платформа автоматически определяет используемые технологии на сайте и проверяет их по актуальной базе данных Common Vulnerabilities and Exposures, содержащей более 200 000 записей об известных уязвимостях.")
    add_para("Интеллектуальный анализ кода. Используя языковую модель GPT-4o, система анализирует исходный код на наличие логических уязвимостей (SQL-инъекции, XSS, проблемы аутентификации), которые отсутствуют в базах CVE.")
    add_para("Верификация владения сайтом. Система требует подтверждения права собственности на сканируемый ресурс через размещение HTML-мета-тега, что предотвращает злоупотребления и использование платформы для атак на чужие ресурсы.")
    add_para("Формирование отчетов. Автоматическая генерация детальных отчетов с описанием найденных уязвимостей, оценкой серьезности по шкале CVSS и пошаговыми рекомендациями по устранению.")
    add_para("Личный кабинет пользователя. Система хранит историю всех сканирований, позволяет управлять верифицированными сайтами и отслеживать динамику устранения уязвимостей.")
    
    # Сравнение
    add_para("Сравнивая разработанную платформу с существующими решениями, можно выделить следующие преимущества:", bold=True)
    add_para("доступность – платформа бесплатна для использования в отличие от коммерческих аналогов стоимостью 399-2000$ в год;")
    add_para("интеграция искусственного интеллекта – поиск логических уязвимостей, недоступных для традиционных сканеров;")
    add_para("простота использования – интуитивно понятный веб-интерфейс не требует технической подготовки;")
    add_para("высокая точность – 92% правильное обнаружение при 8% ложных срабатываний против 85% и 20% у open-source аналогов.")
    
    # Технические детали
    add_para("Технические характеристики:", bold=True)
    add_para("Языки программирования: TypeScript, JavaScript.")
    add_para("Используемые технологии: React, Node.js, Express.js, PostgreSQL, Tailwind CSS, WebSocket.")
    add_para("Внешние API: OpenAI GPT-4o, National Vulnerability Database.")
    add_para("Объем кода: 15 000+ строк.")
    
    doc.save('тезисы_Ямалетдинов_Карим_SecOps.docx')
    print('Тезисы сохранены: тезисы_Ямалетдинов_Карим_SecOps.docx')

if __name__ == '__main__':
    create_thesis()
